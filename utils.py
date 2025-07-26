# utils.py
import PyPDF2, docx2txt

def extract_text(file):
    if file.type == "application/pdf":
        return "".join(PyPDF2.PdfReader(file).pages[i].extract_text() for i in range(len(reader.pages)))
    if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return docx2txt.process(file)
    return file.read().decode("utf-8")

from docx import Document
def write_resume_md(md: str, path="resume.docx"):
    doc = Document()
    for line in md.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        else:
            doc.add_paragraph(line)
    doc.save(path)
